import streamlit as st
from io import BytesIO
from docx import Document
from docx.shared import Mm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ---------------------------------------------------------
# 1. HELPER FUNCTION: GENERATE WORD DOCUMENT
# ---------------------------------------------------------
def create_docx():
    doc = Document()
    
    # Configure Page Size (A4) and Margins
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(20)
    section.right_margin = Mm(15)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'  # Fallback font, systems handle Gujarati differently
    font.size = Pt(11)

    # --- PAGE 1 ---

    # Top Right: Hisabi Patrak
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("હિસાબી પત્રક નંબર ____________")
    
    # Center Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("નવસારી કૃષિ વિશ્વવિધાલય\n")
    run.bold = True
    run.font.size = Pt(16)
    run = p.add_run("મુસાફરી ભથ્થા બીલ")
    run.bold = True
    run.font.size = Pt(14)

    # Info Table (Bill No vs Voucher No)
    table = doc.add_table(rows=1, cols=2)
    table.width = Mm(175)
    table.autofit = False
    
    # Left Column
    cell_l = table.cell(0, 0)
    cell_l.width = Mm(80)
    p = cell_l.paragraphs[0]
    p.add_run("બીલ નંબર :\nતારીખ       :")
    p.runs[0].bold = True

    # Right Column
    cell_r = table.cell(0, 1)
    cell_r.width = Mm(95)
    p = cell_r.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("વાઉચર નં. ____________________\n")
    p.add_run("તારીખ ____________________\n")
    p.add_run("યુનિટ નંબર : __________________\n")
    p.add_run("કોડ નંબર : ____________________")

    doc.add_paragraph() # Spacer

    # Main Body Text
    p = doc.add_paragraph("આચાર્ય અને ડીનશ્રી, નં. મ. કૃષિ મહાવિદ્યાલય, નકૃયું, નવસારી ની કચેરીનું માહે: ઓક્ટોમ્બર - ૨૦૨૫ નું")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    p = doc.add_paragraph("મુસાફરી ભથ્થા બિલ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].underline = True
    p.runs[0].font.size = Pt(14)

    p = doc.add_paragraph("યુનિટ/સબયુનિટ : આચાર્ય અને ડીનશ્રી, ન. મ. કૃષિ મહાવિદ્યાલય, નકૃયું, નવસારી")
    p = doc.add_paragraph("ખર્ચ માટેનું બજેટ સદર :- ____________________")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("યોજનાનું નામ :- ____________________")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph() # Spacer

    # Claim Amount
    p = doc.add_paragraph()
    p.add_run("આથી રૂ।. ")
    r = p.add_run(" 12161 ")
    r.bold = True
    r.underline = True
    p.add_run(" નો દાવો મંજુર કરી ગ્રાહય રાખવામાં આવે છે.")

    # Boxed Amount Area (Simulated with borders)
    tbl_amt = doc.add_table(rows=1, cols=1)
    tbl_amt.style = 'Table Grid'
    cell = tbl_amt.cell(0,0)
    p = cell.paragraphs[0]
    p.add_run("આ બીલમાં જણાવેલ રૂા  ")
    p.add_run("12161").bold = True
    p.add_run("  ( અંકે રૂપિયા ")
    p.add_run("બાર હજાર એકસો એકસઠ પુરા").bold = True
    p.add_run(" પૈસા )\n\n")
    p.add_run("મંજુર કરવામાં આવે છે. અને તે રોકડા / ચેક નં. ______________ તા. ___________ થી ચુકવવામાં આવે છે.")

    doc.add_paragraph() # Spacer

    # Signatures
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.width = Mm(175)
    
    l_cell = sig_table.cell(0, 0)
    l_cell.text = "સ્થળ :    નવસારી\nતારીખ :"
    
    r_cell = sig_table.cell(0, 1)
    p = r_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("____________________\n")
    p.add_run("બીલ મંજુર કરનાર અધિકારીની\nસહી અને હોદ્દો")

    doc.add_paragraph() # Spacer

    # Budget Table
    b_table = doc.add_table(rows=4, cols=3)
    b_table.style = 'Table Grid'
    hdr_cells = b_table.rows[0].cells
    hdr_cells[0].text = ""
    hdr_cells[1].text = "રૂ."
    hdr_cells[2].text = "પૈસા"
    
    b_table.cell(1, 0).text = "(૧) સને ૨૦૨૪-૨૫ માટે બજેટમાં મંજુર થયેલ રકમ"
    b_table.cell(2, 0).text = "(૨) આ બીલ સાથે થયેલ કુલ ખર્ચ"
    b_table.cell(3, 0).text = "(૩) ખર્ચ માટે બાકી રહેતી રકમ"

    doc.add_paragraph() # Spacer

    # Bottom Signature
    p = doc.add_paragraph()
    p.add_run("રૂા ( 12161 ) અંકે રૂપિયા : બાર હજાર એકસો એકસઠ પુરા મંજુર કર્યા")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("નિયંત્રણ અધિકારીની સહી___________").bold = True

    doc.add_page_break()

    # --- PAGE 2 ---

    p = doc.add_paragraph("નોંધ :-")
    p.runs[0].bold = True
    
    notes = [
        "કોલમ નં. ૭ માં મુસાફરી પ્રકાર રેલ્વે/એસ.ટી./હવાઈ/સ્ટીમર/ભાડાનું યુનિવર્સિટી કે સરકારી કે પોતાનું વાહન ઈત્યાદી મારફત કરેલ મુસાફરીની સ્પષ્ટ નોંધ આપવી.",
        "કોલમ નં. ૧૧ થી ૧૩ માઈલેજ મેળવતા અધિકારીઓ કે સભ્યોએ ભરવી.",
        "કોલમ નં. ૧૬ માં માત્ર દૈનિક ભથ્થાની રકમ લેવી જેથી કોલમ (૧૪ X ૧૫= ૧૬) થઈ રહેવું જોઈએ.",
        "જયારે મુસાફરી ભથ્થા બીલમાં શરૂઆતમાં મુસાફરીને બદલે 'હોલ્ટ' દર્શાવવામાં આવેલ હોય તેવા કિસ્સામાં 'હોલ્ટ' ની શરૂઆત થયાની તારીખ કો.નં. ૧૯ માં દર્શાવવી."
    ]
    for note in notes:
        doc.add_paragraph(note, style='List Number')

    doc.add_paragraph()
    p = doc.add_paragraph("યુનિવર્સિટી કર્મચારીએ આપવાનું પ્રમાણપત્ર")
    p.runs[0].bold = True
    
    certs = [
        "આથી પ્રમાણપત્ર આપવામાં આવે છે કે, આ બીલમાં આકારેલ રકમ બીજા કોઈ બીલમાં આકારેલ નથી.",
        "આથી પ્રમાણીત કરવામાં આવે છે કે સદર મુસાફરી ભથ્થા બીલમાં દર્શાવેલ હકીકત સાચી છે...",
        "આથી પ્રમાણપત્ર આપવામાં આવે છે કે બીલમાં દર્શાવેલ પ્રવાસ માટે મેં આ અગાઉ પેશગી લીધેલ નથી...",
        "આ બીલમાં જણાવેલ યુનિવર્સિટી સિવાયની અન્ય સંસ્થાની ભ્રમગીરીના પ્રવાસ માટે...",
        "આથી પ્રમાણપત્ર આપવામાં આવે છે કે, પ્રવાસ ડાયરીમાં દર્શાવવામાં આવેલ સ્થળ, તારીખ, સમય..."
    ]
    for cert in certs:
        doc.add_paragraph(cert, style='List Number')

    # Employee Signature
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("(સચિન આર. પટેલ)\n").bold = True
    p.add_run("સહ પ્રાધ્યાપક\nકર્મચારીની સહી નામ અને હોદ્દો")

    doc.add_paragraph("_" * 50) # Horizontal Line simulation

    # Officer Cert
    p = doc.add_paragraph("યુનિવર્સિટી અધિકારીઓ અને અન્ય સભ્યોએ આપવાનું પ્રમાણપત્ર")
    p.runs[0].bold = True
    doc.add_paragraph("આથી પ્રમાણિત કરવામાં આવે છે કે સદર બીલમાં કરેલ મુસાફરી ભથ્થાનો દાવો આ અંગેના નિયમોની જોગવાઈઓના આધારે ખરો અને યોગ્ય છે.")

    # Officer Signature
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("_______________________\n")
    p.add_run("પ્રાધ્યાપક અને વડા\nકિટકશાત્ર વિભાગ\nનં. મ. કૃષિ મહાવિદ્યાલય\nનકૃયું, નવસારી").bold = True

    doc.add_paragraph()

    # --- FINAL SPLIT TABLE (CALCULATION & RECEIPT) ---
    p = doc.add_paragraph("કર્મચારી / અધિકારી / સભ્યશ્રીએ નીચેની વિગત ભરવી.")
    p.runs[0].bold = True

    final_table = doc.add_table(rows=1, cols=2)
    final_table.style = 'Table Grid'
    final_table.allow_autofit = False
    
    # Left Cell (Calculation)
    c1 = final_table.cell(0, 0)
    c1.width = Mm(90)
    p = c1.paragraphs[0]
    p.add_run("બીલની કુલ રકમ\t= 12161\n")
    p.add_run("બાદ બીલની પેશગીની રકમ\t=\n")
    p.add_run("ચૂકવવા પાત્ર ચોખ્ખી રકમ\t= 12161\n\n")
    p.add_run("પેશગીના નાણાં મળ્યાની તા.\n")
    p.add_run("પેશગી કયા ઝોન /યુનિટમાંથી\t નીલ\nઉપાડવામાં આવી.\n\n")
    p.add_run("પેશગી ઉપાડવાના વાઉચર નંબર ______ તારીખ _____")

    # Right Cell (Receipt)
    c2 = final_table.cell(0, 1)
    c2.width = Mm(90)
    p = c2.paragraphs[0]
    p.add_run("બીલની રકમ રૂ. 12161\n").bold = True
    p.add_run("અંકે રૂપિયા બાર હજાર એકસો એકસઠ પુરા\n").bold = True
    p.add_run("મને મળ્યા છે.\n\n")
    p.add_run("સ્થળ : નવસારી\n")
    p.add_run("તારીખ :\n\n\n")
    p.add_run("                                        (સચિન આર. પટેલ)\n").bold = True
    p.add_run("                                        સહ પ્રાધ્યાપક")

    return doc

# ---------------------------------------------------------
# 2. STREAMLIT APP LOGIC
# ---------------------------------------------------------
st.set_page_config(layout="wide", page_title="Navsari Uni Bill - Final Layout")

# Generate the Word file in memory
doc = create_docx()
buffer = BytesIO()
doc.save(buffer)
buffer.seek(0)

# Sidebar Download Button
st.sidebar.title("Download Options")
st.sidebar.download_button(
    label="Download as Word (.docx)",
    data=buffer,
    file_name="Navsari_Uni_Bill.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

# ---------------------------------------------------------
# 3. CSS for HTML Preview (Existing Code)
# ---------------------------------------------------------
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+Gujarati:wght@400;600;700&display=swap');
    .stApp { background-color: #555; }
    .a4-page {
        background-color: white; color: black; width: 210mm; min-height: 297mm;
        padding: 15mm 20mm; margin: 10px auto;
        font-family: 'Noto Sans Gujarati', sans-serif; font-size: 13px; line-height: 1.5;
        box-shadow: 0 0 15px rgba(0,0,0,0.5);
    }
    .input-line { border-bottom: 1px solid black; display: inline-block; padding: 0 5px; }
    table.budget-table { width: 100%; border-collapse: collapse; margin-top: 10px; font-size: 12px; }
    table.budget-table th, table.budget-table td { border: 1px solid black; padding: 8px; vertical-align: middle; }
    table.budget-table th { text-align: center; font-weight: bold; }
    table.budget-table td { height: 35px; }
    .bold { font-weight: 700; }
    .center { text-align: center; }
    .right { text-align: right; }
    .flex-row { display: flex; justify-content: space-between; align-items: flex-end; }
</style>
""", unsafe_allow_html=True)

# ---------------------------------------------------------
# PAGE 1 PREVIEW HTML
# ---------------------------------------------------------
page1_html = """
<div class="right" style="margin-bottom: 10px; font-size: 12px;">
હિસાબી પત્રક નંબર <span class="input-line" style="width:100px;"></span>
</div>
<div class="center">
<div style="font-size: 22px; font-weight: 700;">નવસારી કૃષિ વિશ્વવિધાલય</div>
<div style="font-size: 16px; font-weight: 700; margin-top: 5px;">મુસાફરી ભથ્થા બીલ</div>
</div>
<div class="flex-row" style="margin-top: 20px;">
<div style="width: 40%; font-weight: 700; font-size: 14px;">
<div style="margin-bottom: 10px;">બીલ નંબર :</div>
<div>તારીખ &nbsp;&nbsp;&nbsp;:</div>
</div>
<div style="width: 50%; text-align: right; font-size: 12px;">
<div style="margin-bottom: 5px;">વાઉચર નં. <span class="input-line" style="width:180px;"></span></div>
<div style="margin-bottom: 5px;">તારીખ <span class="input-line" style="width:180px;"></span></div>
<div style="margin-bottom: 5px;">યુનિટ નંબર : <span class="input-line" style="width:165px;"></span></div>
<div>કોડ નંબર : <span class="input-line" style="width:180px;"></span></div>
</div>
</div>
<div class="center" style="margin-top: 20px; font-weight: 700; font-size: 12px;">
આચાર્ય અને ડીનશ્રી, નં. મ. કૃષિ મહાવિદ્યાલય, નકૃયું, નવસારી ની કચેરીનું માહે: ઓક્ટોમ્બર - ૨૦૨૫ નું
</div>
<div class="center" style="margin-top: 5px;">
<span style="font-size: 16px; font-weight: 700; border-bottom: 2px solid black; padding-bottom: 2px;">મુસાફરી ભથ્થા બિલ</span>
</div>
<div style="margin-top: 20px; font-weight: 600; font-size: 12px;">
યુનિટ/સબયુનિટ : આચાર્ય અને ડીનશ્રી, ન. મ. કૃષિ મહાવિદ્યાલય, નકૃયું, નવસારી
</div>
<div class="center" style="margin-top: 15px; font-weight: 600;">
ખર્ચ માટેનું બજેટ સદર :-
</div>
<div class="center" style="margin-top: 5px; font-weight: 600;">
યોજનાનું નામ :-
</div>
<div style="margin-top: 15px; font-size: 14px;">
આથી રૂા. <span class="input-line" style="width: 150px; text-align:center; font-weight:bold;">12161</span> ના દાવો મંજુર કરી ગ્રાહય રાખવામાં આવે છે.
</div>
<div style="margin-top: 10px; border-top: 1px solid black; padding-top: 10px;">
<div style="line-height: 2;">
આ બીલમાં જણાવેલ રૂા <span class="input-line" style="width: 80px; text-align:center; font-weight:bold;">12161</span> (
<span style="margin-left: 20px;">અંકે રૂપિયા <span style="font-weight:bold; font-size:14px;">બાર હજાર એકસો એકસઠ પુરા</span></span>
<span style="float:right;">પૈસા)</span>
<br>
મંજુર કરવામાં આવે છે. અને તે રોકડા / ચેક નં. <span class="input-line" style="width: 150px;"></span> તા. <span class="input-line" style="width: 100px;"></span> થી ચુકવવામાં આવે છે.
</div>
</div>
<div class="flex-row" style="margin-top: 20px;">
<div>
સ્થળ : &nbsp;&nbsp;&nbsp;નવસારી<br>
તારીખ :
</div>
<div style="text-align: center;">
બીલ મંજુર કરનાર અધિકારીની<br>
સહી અને હોદ્દો
</div>
</div>
<hr style="border-top: 2px solid black; margin: 15px 0;">
<table class="budget-table">
<colgroup>
<col style="width: 70%;">
<col style="width: 15%;">
<col style="width: 15%;">
</colgroup>
<thead>
<tr>
<th rowspan="2" style="text-align:left; border-bottom: none;"></th>
<th>રૂ.</th>
<th>પૈસા</th>
</tr>
<tr>
<th style="height:20px;"></th>
<th style="height:20px;"></th>
</tr>
</thead>
<tbody>
<tr>
<td>(૧) સને ૨૦૨૪-૨૫ માટે બજેટમાં મંજુર થયેલ રકમ</td>
<td></td>
<td></td>
</tr>
<tr>
<td>(૨) આ બીલ સાથે થયેલ કુલ ખર્ચ</td>
<td></td>
<td></td>
</tr>
<tr>
<td>(૩) ખર્ચ માટે બાકી રહેતી રકમ</td>
<td></td>
<td></td>
</tr>
</tbody>
</table>
<div style="margin-top: 15px; font-size: 13px;">
રૂા ( <span class="input-line" style="width: 100px; text-align:center; font-weight:bold;">12161</span> ) અંકે રૂપિયા : <span class="input-line" style="width: 350px; font-weight:bold;">બાર હજાર એકસો એકસઠ પુરા</span> મંજુર કર્યા
</div>
<div style="text-align: right; margin-top: 40px; font-weight: 600;">
નિયંત્રણ અધિકારીની સહી___________
</div>
"""

# ---------------------------------------------------------
# PAGE 2 PREVIEW HTML
# ---------------------------------------------------------
page2_html = """
<div style="font-size: 11px;">
<strong>નોંધ :-</strong>
<ol style="margin-top: 5px; padding-left: 20px; margin-bottom: 10px;">
<li>કોલમ નં. ૭ માં મુસાફરી પ્રકાર રેલ્વે/એસ.ટી./હવાઈ/સ્ટીમર/ભાડાનું યુનિવર્સિટી કે સરકારી કે પોતાનું વાહન ઈત્યાદી મારફત કરેલ મુસાફરીની સ્પષ્ટ નોંધ આપવી.</li>
<li>કોલમ નં. ૧૧ થી ૧૩ માઈલેજ મેળવતા અધિકારીઓ કે સભ્યોએ ભરવી.</li>
<li>કોલમ નં. ૧૬ માં માત્ર દૈનિક ભથ્થાની રકમ લેવી જેથી કોલમ (૧૪ X ૧૫= ૧૬) થઈ રહેવું જોઈએ.</li>
<li>જયારે મુસાફરી ભથ્થા બીલમાં શરૂઆતમાં મુસાફરીને બદલે 'હોલ્ટ' દર્શાવવામાં આવેલ હોય તેવા કિસ્સામાં 'હોલ્ટ' ની શરૂઆત થયાની તારીખ કો.નં. ૧૯ માં દર્શાવવી.</li>
</ol>
</div>
<div style="margin-top: 10px;">
<strong>યુનિવર્સિટી કર્મચારીએ આપવાનું પ્રમાણપત્ર</strong>
<ol style="margin-top: 5px; padding-left: 20px;">
<li>આથી પ્રમાણપત્ર આપવામાં આવે છે કે, આ બીલમાં આકારેલ રકમ બીજા કોઈ બીલમાં આકારેલ નથી.</li>
<li>આથી પ્રમાણીત કરવામાં આવે છે કે સદર મુસાફરી ભથ્થા બીલમાં દર્શાવેલ હકીકત સાચી છે... (નિયમોના ૬૫-૧ ના અનુમાનોને આધારે સાચો છે).</li>
<li>આથી પ્રમાણપત્ર આપવામાં આવે છે કે બીલમાં દર્શાવેલ પ્રવાસ માટે મેં આ અગાઉ પેશગી લીધેલ નથી / <span style="text-decoration: line-through;">મે પેશગી લીધેલ છે...</span></li>
<li>આ બીલમાં જણાવેલ યુનિવર્સિટી સિવાયની અન્ય સંસ્થાની ભ્રમગીરીના પ્રવાસ માટે... (નાણાં મળશે તો જમા કરાવવામાં આવશે).</li>
<li>આથી પ્રમાણપત્ર આપવામાં આવે છે કે, પ્રવાસ ડાયરીમાં દર્શાવવામાં આવેલ સ્થળ, તારીખ, સમય, કિલોમીટર કચેરીના વાહન લોગબુક મુજબ આકારવામાં આવેલ છે.</li>
</ol>
</div>
<div class="flex-row" style="margin-top: 20px;">
<div></div>
<div style="text-align: center;">
<strong>(સચિન આર. પટેલ)</strong><br>
સહ પ્રાધ્યાપક<br>
કર્મચારીની સહી નામ અને હોદ્દો
</div>
</div>
<div class="center" style="margin-top: 30px; font-weight: 700; font-size: 14px;">
યુનિવર્સિટી અધિકારીઓ અને અન્ય સભ્યોએ આપવાનું પ્રમાણપત્ર
</div>
<div class="center" style="margin-top: 10px; font-size: 12px; font-weight: 600;">
આથી પ્રમાણિત કરવામાં આવે છે કે સદર બીલમાં કરેલ મુસાફરી ભથ્થાનો દાવો આ અંગેના નિયમોની જોગવાઈઓના આધારે ખરો અને યોગ્ય છે.
</div>
<div style="margin-top: 20px; display: flex; justify-content: flex-end;">
<div style="text-align: center; font-weight: 700; font-size: 13px; line-height: 1.4;">
પ્રાધ્યાપક અને વડા<br>
કિટકશાત્ર વિભાગ<br>
નં. મ. કૃષિ મહાવિદ્યાલય<br>
નકૃયું, નવસારી
</div>
</div>
<div style="margin-top: 10px; font-weight: 700; font-size: 13px; border-bottom: 2px solid black; padding-bottom: 2px;">
કર્મચારી / અધિકારી / સભ્યશ્રીએ નીચેની વિગત ભરવી.
</div>
<div style="display: flex; border-bottom: 2px solid black;">
<div style="width: 50%; border-right: 2px solid black; padding: 10px 10px 10px 0; font-weight: 600;">
<div style="display: flex; justify-content: space-between; margin-bottom: 5px;">
<span>બીલની કુલ રકમ</span>
<span style="font-weight: 700; font-size: 14px;">= &nbsp;&nbsp; 12161</span>
</div>
<div style="display: flex; justify-content: space-between; margin-bottom: 5px;">
<span>બાદ બીલની પેશગીની રકમ</span>
<span style="font-weight: 700;">=</span>
</div>
<div style="display: flex; justify-content: space-between; margin-bottom: 15px;">
<span>ચૂકવવા પાત્ર ચોખ્ખી રકમ</span>
<span style="font-weight: 700; font-size: 14px;">= &nbsp;&nbsp; 12161</span>
</div>
<div style="margin-bottom: 5px;">પેશગીના નાણાં મળ્યાની તા.</div>
<div style="margin-bottom: 5px; display: flex;">
<span>પેશગી કયા ઝોન /યુનિટમાંથી<br>ઉપાડવામાં આવી.</span>
<span style="margin-left: auto; align-self: center; font-weight: bold; margin-right: 20px;">નીલ</span>
</div>
<div style="margin-top: 15px;">
પેશગી ઉપાડવાના વાઉચર નંબર<br>
<div style="text-align: center; border-bottom: 1px solid black; width: 100px; margin-top: 5px;">&nbsp;</div>
તારીખ
</div>
</div>
<div style="width: 50%; padding: 10px 0 10px 10px; position: relative; font-weight: 600;">
<div style="margin-bottom: 5px;">
બીલની રકમ રૂ. <span style="font-weight: 700; border-bottom: 1px solid black; padding: 0 10px; font-size: 14px;">12161</span>
</div>
<div style="margin-bottom: 15px;">
અંકે રૂપિયા <span style="font-weight: 700; border-bottom: 1px solid black; font-size: 14px;">બાર હજાર એકસો એકસઠ પુરા</span>
</div>
<div style="margin-bottom: 20px;">મને મળ્યા છે.</div>
<div style="margin-bottom: 5px;">
સ્થળ : <span style="border-bottom: 1px solid black;">નવસારી</span>
</div>
<div style="margin-bottom: 50px;">
તારીખ :
</div>
<div style="text-align: center; position: absolute; bottom: 10px; right: 10px;">
<div style="font-weight: 700;">(સચિન આર. પટેલ)</div>
<div style="font-weight: 600;">સહ પ્રાધ્યાપક</div>
<div style="font-weight: 600;">કર્મચારીની સહી</div>
</div>
</div>
</div>
"""

# ---------------------------------------------------------
# RENDER COLUMNS
# ---------------------------------------------------------
col1, col2 = st.columns(2)
with col1:
    st.markdown(f'<div class="a4-page">{page1_html}</div>', unsafe_allow_html=True)
with col2:
    st.markdown(f'<div class="a4-page">{page2_html}</div>', unsafe_allow_html=True)
